import os
import streamlit as st
import time
from langchain_groq import ChatGroq
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain.chains.combine_documents import create_stuff_documents_chain
from langchain_core.prompts import ChatPromptTemplate
from langchain.chains import create_retrieval_chain
from langchain_community.vectorstores import FAISS
from langchain_community.document_loaders import PyPDFLoader
from langchain_google_genai import GoogleGenerativeAIEmbeddings
from docx import Document
from langchain.schema import Document as LangchainDocument
import warnings
import shutil

st.set_page_config(page_title="DocuQuery AI", page_icon=":robot_face:")
# Suppress specific warning
warnings.filterwarnings("ignore", message="Importing verbose from langchain root module is no longer supported.")

# Load the GROQ And Google API KEY
groq_api_key = st.secrets['GROQ_API_KEY']
os.environ["google_api_key"] = st.secrets['GOOGLE_API_KEY']

st.title("DocuQuery AI :robot_face:")

models = {
    "gemma-7b-it": "gemma-7b-it",
    "gemma2-9b-it": "gemma2-9b-it",
    "mixtral-8x7b-32768": "Mixtral-8x7b-32768",
    "llama3-70b-8192": "llama3-70b-8192",
    "llama-3.1-70b-versatile": "llama-3.1-70b-versatile",
    "llama3-8B": "llama3-8b-8192",
    "llama-3.1-8b-instant": "llama-3.1-8b-instant",
    "llama3-groq-8b-8192-tool-use-preview": "llama3-groq-8b-8192-tool-use-preview",
    "llama3-groq-70b-8192-tool-use-preview": "llama3-groq-70b-8192-tool-use-preview"
}

selected_model = st.sidebar.selectbox(
    "**Select Model**",
    list(models.keys())
)

llm = ChatGroq(groq_api_key=groq_api_key,
               model_name=models[selected_model])

prompt_temp = ChatPromptTemplate.from_template(
    """
    Answer the questions based on the provided context only.
    Please provide the most accurate response based on the question
    <context>
    {context}
    <context>
    Questions:{input}
    """
)

# Create a temporary directory to save uploaded files
temp_dir = "temp_files"
if not os.path.exists(temp_dir):
    os.makedirs(temp_dir)


def load_docx(file_path):
    """Function to load and extract text from a DOCX file"""
    doc = Document(file_path)
    full_text = []
    for para in doc.paragraphs:
        full_text.append(para.text)
    return '\n'.join(full_text)


def vector_embedding(uploaded_files):
    if "vectors" not in st.session_state:
        st.session_state.embeddings = GoogleGenerativeAIEmbeddings(model="models/embedding-001")
        documents = []

        # Process each uploaded file
        for uploaded_file in uploaded_files:
            # Save the uploaded file temporarily
            temp_file_path = os.path.join(temp_dir, uploaded_file.name)
            with open(temp_file_path, "wb") as f:
                f.write(uploaded_file.getbuffer())

            # Load the file content based on the file type
            if uploaded_file.name.endswith('.pdf'):
                loader = PyPDFLoader(temp_file_path)
                loaded_docs = loader.load()
                for doc in loaded_docs:
                    documents.append(LangchainDocument(page_content=doc.page_content, metadata=doc.metadata))
            elif uploaded_file.name.endswith('.docx'):
                docx_text = load_docx(temp_file_path)
                documents.append(LangchainDocument(page_content=docx_text, metadata={'source': temp_file_path}))

        st.session_state.text_splitter = RecursiveCharacterTextSplitter(chunk_size=1000, chunk_overlap=200)
        st.session_state.final_documents = st.session_state.text_splitter.split_documents(documents)
        st.session_state.vectors = FAISS.from_documents(st.session_state.final_documents, st.session_state.embeddings)

        # Delete the temporary files after vector creation
        for uploaded_file in uploaded_files:
            temp_file_path = os.path.join(temp_dir, uploaded_file.name)
            os.remove(temp_file_path)


uploaded_files = st.sidebar.file_uploader("Upload PDF or DOCX files", type=["pdf", "docx"], accept_multiple_files=True)

if st.sidebar.button("Upload & Process Files"):
    if uploaded_files:
        vector_embedding(uploaded_files)
        st.write(
            "Your uploaded files are converted into Vector Store DB. Now it is ready. You can ask any question from the uploaded files.")
    else:
        st.error("Please upload PDF or DOCX files first.")

prompt = st.chat_input("ask anything!")

if prompt:
    if "vectors" in st.session_state:
        document_chain = create_stuff_documents_chain(llm, prompt_temp)
        retriever = st.session_state.vectors.as_retriever()
        retrieval_chain = create_retrieval_chain(retriever, document_chain)
        start = time.process_time()
        response = retrieval_chain.invoke({'input': prompt})
        st.divider()
        col1, col2 = st.columns(2)
        with col1:
            st.write(f"Response time: {time.process_time() - start}")
        with col2:
            st.write(f"Selected Model: {selected_model}")

        st.success(response['answer'])

        # With a streamlit expander
        with st.expander("Document Similarity Search"):
            # Find the relevant chunks
            for i, doc in enumerate(response["context"]):
                st.write(doc.page_content)
                st.write("--------------------------------")
    else:
        st.error("Please upload and process the documents first! open sidebar! 👈")

st.sidebar.info(
    "Welcome to DocuQuery AI! This application allows you to upload and analyze documents for quick information retrieval.\n\n"
    "How to use:\n"
    "1. Upload your PDF or DOCX files using the uploader below.\n"
    "2. Click the 'Upload & Process Files' button to process your files.\n"
    "3. Once processed, you can enter any question related to the content of the uploaded documents, and the AI will provide an accurate response.\n\n"
    "Please follow the steps to get started."
)

st.sidebar.divider()
st.sidebar.write(
    "Developed🚀 by **Rakesh Kumar**\n"
    "\n Feel free to connect and share your feedback(✨)"
    "\n on LinkedIn: [Rakesh Kumar](https://www.linkedin.com/in/m-rakesh-kr/)"
)
st.sidebar.divider()

st.sidebar.text('© 2024 DocuQuery AI.')


# Cleanup temporary directory on session state initialization
def cleanup_temp_dir():
    if os.path.exists(temp_dir):
        shutil.rmtree(temp_dir)
        os.makedirs(temp_dir)


if "initialized" not in st.session_state:
    st.session_state.initialized = True
    cleanup_temp_dir()
